#!/usr/bin/env python3
"""Generate an editable Word resume for Tanvi Patel."""

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt, RGBColor


NAVY = RGBColor(0x14, 0x23, 0x3A)


def set_run_font(run, name="Calibri", size=10.5, bold=False, italic=False, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color


def set_paragraph_spacing(paragraph, before=0, after=4, line=1.15):
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE


def add_bottom_border(paragraph):
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "14233A")
    pBdr.append(bottom)
    pPr.append(pBdr)


def heading(doc, text):
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=10, after=4, line=1.0)
    run = p.add_run(text.upper())
    set_run_font(run, size=12, bold=True, color=NAVY)
    add_bottom_border(p)
    return p


def job_header(doc, title, dates):
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=6, after=0, line=1.05)
    tab_stops = p.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Inches(7.0))
    run = p.add_run(title)
    set_run_font(run, size=11, bold=True, color=NAVY)
    p.add_run("\t")
    run = p.add_run(dates)
    set_run_font(run, size=10.5, bold=True, color=NAVY)
    return p


def org_line(doc, text):
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=0, after=2, line=1.05)
    run = p.add_run(text)
    set_run_font(run, size=10.5, italic=True)
    return p


def bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        set_paragraph_spacing(p, before=0, after=1, line=1.12)
        p.clear()
        run = p.add_run(item)
        set_run_font(run, size=10.5)
        p.paragraph_format.left_indent = Inches(0.25)


def note(doc, label, text):
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=1, after=2, line=1.12)
    run = p.add_run(f"{label}: ")
    set_run_font(run, size=10.5, bold=True)
    run = p.add_run(text)
    set_run_font(run, size=10.5)


def skill(doc, label, text):
    p = doc.add_paragraph(style="List Bullet")
    set_paragraph_spacing(p, before=0, after=1, line=1.12)
    p.clear()
    run = p.add_run(f"{label}: ")
    set_run_font(run, size=10.5, bold=True, color=NAVY)
    run = p.add_run(text)
    set_run_font(run, size=10.5)
    p.paragraph_format.left_indent = Inches(0.25)


def main():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)

    name = doc.add_paragraph()
    name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(name, before=0, after=2, line=1.0)
    run = name.add_run("Tanvi Patel")
    set_run_font(run, size=22, bold=True, color=NAVY)

    contact = doc.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(contact, before=0, after=2, line=1.15)
    run = contact.add_run(
        "Email: tanvipatel2596@gmail.com  |  Contact: +61 431 982 502  |  "
        "LinkedIn: tanvi-patel-ua3013  |  NSW Driver Licence: Class C"
    )
    set_run_font(run, size=10)

    heading(doc, "About me")
    about = doc.add_paragraph()
    set_paragraph_spacing(about, before=2, after=4, line=1.15)
    about.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = about.add_run(
        "Results-driven Digital Marketing Coordinator and Design Specialist with 5+ years of experience in "
        "social media strategy, content creation, and brand marketing. Currently leading marketing for "
        "Global Medics and Medacs Healthcare at MGG Health, owning social media end to end across LinkedIn, "
        "Instagram, and Facebook, plus events and merchandise. Background in performance marketing across "
        "Meta and TikTok, with experience briefing and managing specialists who run paid media. Skilled in "
        "campaigns that lift engagement, generate leads, and strengthen brand identity. Proficient in Adobe "
        "Creative Suite, video production, and marketing analytics."
    )
    set_run_font(run, size=10.5)

    heading(doc, "Skills")
    skill(doc, "Social Media Management", "End-to-end LinkedIn, Instagram, and Facebook — strategy, content, publishing, community management, and reporting.")
    skill(doc, "Social Media Advertising", "Experienced in Meta Ads Manager, Instagram, and TikTok Ads for lead generation and ROI. Able to brief and manage people handling paid media.")
    skill(doc, "Events & Merchandise", "Plan and deliver brand events and merchandise programmes that support visibility and stakeholder engagement.")
    skill(doc, "Campaign Strategy & Content", "Full-funnel campaigns and branded assets (ads, reels, posts, landing pages) aligned to voice and conversion goals. Skilled in short-form video with Premiere Pro, After Effects, and TikTok tools.")
    skill(doc, "Influencer & PR Management", "Identify, negotiate, and manage influencer partnerships to boost brand visibility and engagement.")
    skill(doc, "Analytics & CRM", "Google Analytics, Meta Insights, LinkedIn Analytics, Shopify, Klaviyo, and Dreamweaver for reporting, automations, and newsletters. Familiar with SEO/SEM.")
    skill(doc, "Collaboration", "Work cross-functionally with sales, clinical, and creative teams, and manage specialists delivering paid media.")

    heading(doc, "Professional Work Experience")

    job_header(doc, "Marketing Executive", "03/26 – Present")
    org_line(doc, "MGG Health — Global Medics & Medacs Healthcare, Sydney")
    bullets(doc, [
        "Own end-to-end social media for Global Medics and Medacs Healthcare across LinkedIn, Instagram, and Facebook — strategy, content, publishing, community management, and reporting.",
        "Lead marketing for two healthcare staffing brands within the MGG Health group, aligning activity with candidate attraction and client brand positioning.",
        "Plan and deliver events and merchandise programmes that support brand presence, candidate engagement, and client relationships.",
        "Brief and manage paid media specialists as needed, providing brand direction and performance oversight without hands-on media buying.",
    ])

    job_header(doc, "Social Media and Marketing Executive", "03/25 – 01/26")
    org_line(doc, "Brightstar Nursing Australia, Connecting Care, SuccessVisa, Sydney")
    bullets(doc, [
        "Lead digital campaigns and paid social promotions to boost brand awareness and lead generation.",
        "Collaborated with agents and creators to produce lifestyle-oriented content that resonated with audiences.",
        "Manage Meta and Instagram ads, optimising for audience engagement and cost per conversion.",
        "Create video content using Premiere Pro, After Effects, and TikTok, improving reach and retention.",
        "Develop branded visuals and marketing materials (flyers, ads, email banners) aligned with campaign objectives.",
        "Analyse campaign performance and prepare reports to guide strategy.",
    ])

    job_header(doc, "Graphic Designer and Social Media Marketing", "12/24 – 03/25")
    org_line(doc, "Wiseberry Real Estate, Sydney")
    bullets(doc, [
        "Manage 3+ social media pages, creating engaging content, posts, and Meta ads using Illustrator and Photoshop.",
        "Produce video content using Premiere Pro and After Effects, collaborating with creators.",
        "Design DL flyers, magazines (InDesign), and marketing materials; send emails via CRM, Amazon S3, and Dreamweaver.",
        "Developed knowledge of Jira and Meistertask.",
    ])
    note(doc, "Software used", "Premiere Pro, After Effects, Illustrator, Photoshop, InDesign, Dreamweaver, Amazon S3, Jira, Meistertask, Figma")
    note(doc, "Achievement", "Boosted engagement across multiple social pages with creative content, video, email campaigns, and ads that drove customer interaction and retention.")

    job_header(doc, "Graduate Planner / Marketing", "03/24 – 12/24")
    org_line(doc, "Macroplan, Sydney")
    bullets(doc, [
        "Developed digital and print materials, supporting project phases with strong visual content.",
        "Managed end-to-end design for presentations, brochures, flyers, and client meeting slides.",
        "Converted architectural and planning diagrams into digital format; devised architectural solutions for project planning.",
    ])
    note(doc, "Software used", "Adobe Illustrator, Adobe Photoshop, Adobe InDesign, AutoCAD, ArcGIS")

    job_header(doc, "Graphic Designer & Marketing Associate", "03/23 – 02/24")
    org_line(doc, "IBD Medical, Sydney")
    bullets(doc, [
        "Produced video and graphic assets that enhanced brand storytelling and audience engagement.",
        "Designed educational visuals and published website articles on mental health and diabetes management.",
        "Collaborated with stakeholders to implement marketing strategy; created Klaviyo campaigns and analysed Shopify user behaviour.",
        "Maintained brand consistency using Photoshop, Illustrator, Canva, and Procreate; developed knowledge of Trello.",
    ])
    note(doc, "Software used", "Illustrator, InDesign, Photoshop, Shopify, Klaviyo, Canva")
    note(doc, "Achievement", "Marketing campaigns increased customer inflow and strengthened IBD Medical brand value.")

    job_header(doc, "Architect / Graphic Designer", "08/20 – 09/21")
    org_line(doc, "HCP Design, Planning and Management Pvt. Ltd, Ahmedabad, India")
    bullets(doc, [
        "Translated the company's ideology and projects into visuals, including charts, diagrams, and illustrations.",
        "Researched and designed maps for over 40 Indian cities, highlighting key urban features.",
    ])
    note(doc, "Software used", "AutoCAD, ArcGIS, Illustrator, Photoshop and InDesign")
    note(doc, "Achievement", 'Pivotal role in creating the "Urban Design and Planning" book; analysed urban plans of 40+ Indian cities.')

    job_header(doc, "Architectural Intern / Graphic Designer", "12/17 – 06/18")
    org_line(doc, "Artha Studio, Pune, India")
    bullets(doc, [
        "Led a furniture documentation team covering photography and furniture drawings; took site measurements for planning stages.",
        "Researched materials and structures while collaborating with vendors to find solutions.",
    ])

    heading(doc, "Education")
    job_header(doc, "Master of Urban and Regional Planning", "2022 – 2023")
    org_line(doc, "The University of Sydney, Sydney, Australia")
    job_header(doc, "Bachelor of Architecture", "2014 – 2019")
    org_line(doc, "CEPT University, Ahmedabad, India")

    out = "/workspace/resume/Tanvi_Patel_Resume.docx"
    doc.save(out)
    print(f"Wrote {out}")


if __name__ == "__main__":
    main()
